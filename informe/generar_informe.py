#!/usr/bin/env python3
"""Genera el informe 'Uso de ProjectLibre en el Proyecto Lush Nails SPA'
con normas APA 7ª edición usando python-docx.

Uso:
    python3 generar_informe.py            # genera informe_apa7.docx
    python3 generar_informe.py -o nombre  # genera con otro nombre

Las capturas se leen de ./capturas/ (ver diccionario CAPTURAS). Si un archivo
no existe, se dibuja un recuadro de espacio para pegar la captura.
"""

import argparse
import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Configuración del usuario
# ---------------------------------------------------------------------------
DATOS = {
    "titulo": "Uso de ProjectLibre en la Gestión de Proyectos de Software",
    "subtitulo": "Aplicación al Portal Empresarial Lush Nails SPA",
    "autor": "Nombre del Estudiante",
    "afiliacion": "Nombre de la Institución",
    "curso": "Gestión de Proyectos de Software",
    "docente": "Nombre del Docente",
    "fecha": "14 de agosto de 2026",
}

# Cada entrada: nombre del archivo en ./capturas/ -> descripción de la figura.
CAPTURAS = {
    "01_instalacion.png": "Ventana de bienvenida de ProjectLibre tras la instalación",
    "02_nuevo_proyecto.png": "Diálogo 'Nuevo proyecto' para definir nombre, fechas y horario laboral",
    "03_estructura_tareas.png": "Estructura de Descomposición del Trabajo (WBS) del proyecto Lush Nails",
    "04_dependencias.png": "Asignación de dependencias entre tareas (fin comienzo, retrasos)",
    "05_recursos.png": "Hoja de recursos con el equipo asignado al proyecto",
    "06_asignacion.png": "Asignación de recursos a cada tarea del cronograma",
    "07_gantt.png": "Diagrama de Gantt con el cronograma completo del proyecto",
    "08_ruta_critica.png": "Resaltado de la ruta crítica del proyecto",
    "09_costos.png": "Vista de costos por tarea y costo total del proyecto",
    "10_guardar.png": "Archivo .pod del proyecto guardado para su seguimiento",
}

# Tareas reales del proyecto Lush Nails (nombre, duración, dependencia, recurso)
CRONOGRAMA = [
    ("1.1", "Levantar base de datos PostgreSQL (Docker)", 1, "-", "Desarrollador BD"),
    ("1.2", "Diseño del modelo entidad-relación (19 tablas)", 3, "-", "Desarrollador BD"),
    ("1.3", "Definición de vistas financieras (v_financiero, v_citas_completas)", 2, "1.2", "Desarrollador BD"),
    ("2.1", "Panel admin: autenticación y roles (Express + EJS)", 4, "1.1", "Backend"),
    ("2.2", "Módulo de citas y clientes", 3, "2.1", "Backend"),
    ("2.3", "Módulo de historial de atenciones", 2, "2.2", "Backend"),
    ("2.4", "Balanced Scorecard y tablero de comando", 3, "2.3", "Backend"),
    ("3.1", "Sitio web corporativo (React 19)", 5, "2.1", "Frontend"),
    ("3.2", "Integración web con API del panel", 3, "3.1", "Frontend"),
    ("4.1", "Diagramas BPMN de procesos", 2, "1.2", "Analista"),
    ("4.2", "Pruebas integrales y revisión de entregables", 3, "2.4;3.2;4.1", "Analista"),
    ("4.3", "Presentación final del proyecto", 1, "4.2", "Analista"),
]

REFERENCIAS = [
    "ProjectLibre. (2024). ProjectLibre: Software libre de gestión de proyectos. https://www.projectlibre.org/",
    "Project Management Institute. (2021). A guide to the project management body of knowledge (PMBOK guide) (7.ª ed.). Project Management Institute.",
    "American Psychological Association. (2020). Publication manual of the American Psychological Association (7.ª ed.). https://doi.org/10.1037/0000165-000",
]


# ---------------------------------------------------------------------------
# Utilidades APA
# ---------------------------------------------------------------------------
def configurar_documento(doc):
    for seccion in doc.sections:
        seccion.top_margin = Cm(2.54)
        seccion.bottom_margin = Cm(2.54)
        seccion.left_margin = Cm(2.54)
        seccion.right_margin = Cm(2.54)

    estilo = doc.styles["Normal"]
    estilo.font.name = "Times New Roman"
    estilo.font.size = Pt(12)
    estilo._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = estilo.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)


def agregar_numero_pagina(doc):
    seccion = doc.sections[0]
    header = seccion.header
    parrafo = header.paragraphs[0]
    parrafo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    corrida = parrafo.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    corrida._r.append(fldChar1)
    corrida._r.append(instrText)
    corrida._r.append(fldChar2)
    corrida.font.name = "Times New Roman"
    corrida.font.size = Pt(12)


def pagina_titulo(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(DATOS["titulo"])
    r.bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(DATOS["subtitulo"]).bold = True
    for _ in range(4):
        doc.add_paragraph()
    for campo in ("autor", "afiliacion", "curso", "docente", "fecha"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(DATOS[campo])
    doc.add_page_break()


def parrafo_apa(doc, texto, alineacion=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph(texto)
    p.alignment = alineacion
    p.paragraph_format.first_line_indent = Inches(0.5)
    return p


def encabezado1(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(texto)
    r.bold = True
    p.paragraph_format.space_before = Pt(12)
    return p


def encabezado2(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(texto)
    r.bold = True
    return p


def nuevo_pagina_encabezado1(doc, texto):
    doc.add_page_break()
    encabezado1(doc, texto)


def figura(doc, descripcion, numero, archivo):
    """Inserta una imagen si existe; si no, un recuadro para pegarla."""
    ruta = os.path.join(os.path.dirname(os.path.abspath(__file__)), "capturas", archivo)
    if os.path.exists(ruta):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(ruta, width=Inches(5.5))
    else:
        tabla = doc.add_table(rows=1, cols=1)
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
        celda = tabla.cell(0, 0)
        celda.text = ""
        p = celda.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("PEGAR CAPTURA AQUÍ")
        r.bold = True
        r.font.color.rgb = RGBColor(150, 150, 150)
        # bordes
        tcPr = celda._tc.get_or_add_tcPr()
        bordes = OxmlElement("w:tcBorders")
        for lado in ("top", "left", "bottom", "right"):
            b = OxmlElement(f"w:{lado}")
            b.set(qn("w:val"), "dashed")
            b.set(qn("w:sz"), "8")
            b.set(qn("w:color"), "888888")
            bordes.append(b)
        tcPr.append(bordes)
        celda.height = Cm(5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    r = p.add_run(f"Figura {numero}")
    r.bold = True
    r.italic = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.first_line_indent = Inches(0)
    r2 = p2.add_run(f"Nota. {descripcion}.")
    r2.italic = True
    r2.font.size = Pt(11)


def tabla_cronograma(doc):
    doc.add_paragraph()
    filas = len(CRONOGRAMA) + 1
    tabla = doc.add_table(rows=filas, cols=5)
    tabla.style = "Table Grid"
    encabezados = ["ID", "Tarea", "Duración (días)", "Dependencia", "Recurso"]
    for i, enc in enumerate(encabezados):
        celda = tabla.cell(0, i)
        celda.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        celda.paragraphs[0].add_run(enc).bold = True
    for f, (tid, nombre, dur, dep, rec) in enumerate(CRONOGRAMA, start=1):
        datos = (tid, nombre, str(dur), dep, rec)
        for c, val in enumerate(datos):
            celda = tabla.cell(f, c)
            celda.paragraphs[0].add_run(val)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    r = p.add_run("Tabla 1")
    r.bold = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.first_line_indent = Inches(0)
    p2.add_run("Nota. Cronograma propuesto del proyecto en ProjectLibre.").italic = True
    p2.runs[0].font.size = Pt(11)


def referencias_apa(doc):
    nuevo_pagina_encabezado1(doc, "Referencias")
    for ref in REFERENCIAS:
        p = doc.add_paragraph(ref)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.left_indent = Inches(0.5)


# ---------------------------------------------------------------------------
# Cuerpo del informe
# ---------------------------------------------------------------------------
def cuerpo(doc):
    # ---------------- Introducción ----------------
    nuevo_pagina_encabezado1(doc, "Introducción")
    parrafo_apa(doc, DATOS["intro"])

    # ---------------- Descripción del proyecto ----------------
    encabezado2(doc, "Descripción del Proyecto")
    parrafo_apa(doc, DATOS["desc_proyecto"])
    parrafo_apa(doc, DATOS["desc_proyecto2"])

    # ---------------- ProjectLibre ----------------
    nuevo_pagina_encabezado1(doc, "ProjectLibre como Herramienta de Gestión")
    parrafo_apa(doc, DATOS["que_es"])

    # ---------------- Instalación ----------------
    encabezado2(doc, "Instalación de ProjectLibre en Linux")
    parrafo_apa(doc, DATOS["instalacion"])
    figura(doc, CAPTURAS["01_instalacion.png"], 1, "01_instalacion.png")

    # ---------------- Crear proyecto ----------------
    encabezado2(doc, "Creación de un Nuevo Proyecto")
    parrafo_apa(doc, DATOS["nuevo_proyecto"])
    figura(doc, CAPTURAS["02_nuevo_proyecto.png"], 2, "02_nuevo_proyecto.png")

    # ---------------- WBS ----------------
    encabezado2(doc, "Definición de Tareas (Estructura de Descomposición del Trabajo)")
    parrafo_apa(doc, DATOS["wbs"])
    tabla_cronograma(doc)
    figura(doc, CAPTURAS["03_estructura_tareas.png"], 3, "03_estructura_tareas.png")

    # ---------------- Dependencias ----------------
    encabezado2(doc, "Dependencias entre Tareas")
    parrafo_apa(doc, DATOS["dependencias"])
    figura(doc, CAPTURAS["04_dependencias.png"], 4, "04_dependencias.png")

    # ---------------- Recursos ----------------
    encabezado2(doc, "Recursos del Proyecto")
    parrafo_apa(doc, DATOS["recursos"])
    figura(doc, CAPTURAS["05_recursos.png"], 5, "05_recursos.png")
    figura(doc, CAPTURAS["06_asignacion.png"], 6, "06_asignacion.png")

    # ---------------- Gantt ----------------
    encabezado2(doc, "Diagrama de Gantt")
    parrafo_apa(doc, DATOS["gantt"])
    figura(doc, CAPTURAS["07_gantt.png"], 7, "07_gantt.png")

    # ---------------- Ruta crítica ----------------
    encabezado2(doc, "Ruta Crítica")
    parrafo_apa(doc, DATOS["ruta_critica"])
    figura(doc, CAPTURAS["08_ruta_critica.png"], 8, "08_ruta_critica.png")

    # ---------------- Costos ----------------
    encabezado2(doc, "Costos del Proyecto")
    parrafo_apa(doc, DATOS["costos"])
    figura(doc, CAPTURAS["09_costos.png"], 9, "09_costos.png")

    # ---------------- Seguimiento ----------------
    encabezado2(doc, "Guardado y Seguimiento del Proyecto")
    parrafo_apa(doc, DATOS["guardado"])
    figura(doc, CAPTURAS["10_guardar.png"], 10, "10_guardar.png")

    # ---------------- Aplicación al proyecto ----------------
    nuevo_pagina_encabezado1(doc, "Aplicación de ProjectLibre al Proyecto Lush Nails")
    parrafo_apa(doc, DATOS["aplicacion"])

    # ---------------- Conclusiones ----------------
    nuevo_pagina_encabezado1(doc, "Conclusiones")
    parrafo_apa(doc, DATOS["conclusiones"])


# ---------------------------------------------------------------------------
# Contenido (texto redactado)
# ---------------------------------------------------------------------------
DATOS.update(
    {
        "intro": (
            "La gestión de proyectos de software requiere de herramientas que permitan planificar, "
            "organizar y dar seguimiento a las actividades, los recursos y los tiempos involucrados en el "
            "desarrollo de un sistema de información. ProjectLibre es un software libre de código abierto "
            "que brinda las funcionalidades esenciales para la administración de proyectos, ofreciendo una "
            "alternativa de uso gratuito compatible con los estándares de la gestión profesional."
        ),
        "desc_proyecto": (
            "El presente informe describe la utilización de ProjectLibre como herramienta de planificación "
            "aplicada al proyecto denominado Portal Empresarial Lush Nails SPA, un sistema de gestión de "
            "citas, servicios, clientes y panel administrativo para un spa de belleza. El proyecto está "
            "compuesto por tres componentes principales: una base de datos PostgreSQL, un panel de "
            "administración desarrollado con Node.js, Express y EJS, y un sitio web corporativo desarrollado "
            "con React."
        ),
        "desc_proyecto2": (
            "Para el desarrollo del proyecto se definió un cronograma de trabajo en el que se identifican las "
            "tareas de diseño de la base de datos, desarrollo del panel administrativo, desarrollo del sitio "
            "web, documentación y pruebas, así como los recursos humanos necesarios para cada etapa."
        ),
        "que_es": (
            "ProjectLibre es una aplicación de gestión de proyectos de código abierto, escrita en Java, que "
            "funciona como alternativa gratuita a Microsoft Project. Permite crear cronogramas, administrar "
            "tareas, asignar recursos, calcular costos y visualizar el diagrama de Gantt, así como identificar "
            "la ruta crítica del proyecto. Es multiplataforma y se encuentra disponible para Windows, macOS y "
            "Linux (ProjectLibre, 2024)."
        ),
        "instalacion": (
            "En el sistema operativo Linux (CachyOS, basado en Arch), ProjectLibre se instala desde el "
            "repositorio de usuarios de Arch (AUR) mediante el comando paru -S projectlibre. Durante la "
            "instalación se instalan automáticamente las dependencias requeridas, entre las que destaca el "
            "entorno de ejecución de Java (OpenJDK 21). Una vez finalizada la instalación, la aplicación se "
            "abre desde el menú de aplicaciones o con el comando projectlibre en la terminal."
        ),
        "nuevo_proyecto": (
            "Al iniciar ProjectLibre se muestra la ventana de bienvenida. Para comenzar se selecciona la "
            "opción Nuevo proyecto, se asigna un nombre al proyecto y se definen las fechas de inicio y fin, "
            "así como el calendario laboral, considerando días hábiles y feriados. Para el proyecto Lush Nails "
            "se definió un calendario de lunes a viernes con jornada de ocho horas."
        ),
        "wbs": (
            "La Estructura de Descomposición del Trabajo (WBS) permite organizar el proyecto en tareas "
            "jerárquicas. En la Tabla 1 se presenta el cronograma propuesto para el proyecto Lush Nails, en el "
            "que se identifican cuatro fases principales: base de datos, panel administrativo, sitio web y "
            "documentación con pruebas."
        ),
        "dependencias": (
            "Las dependencias establecen el orden lógico de ejecución de las tareas. En ProjectLibre se "
            "definen mediante la relación entre una tarea predecesora y una tarea sucesora, con los tipos "
            "fin-comienzo (la más común), comienzo-comienzo, fin-fin y comienzo-fin. Por ejemplo, el desarrollo "
            "del módulo de citas depende de que la autenticación y los roles estén implementados."
        ),
        "recursos": (
            "En la hoja de recursos se registran los recursos humanos del proyecto: desarrollador de base de "
            "datos, desarrollador backend, desarrollador frontend y analista. A cada recurso se le asigna una "
            "tarifa horaria o diaria, lo que permite calcular de forma automática el costo de cada tarea."
        ),
        "gantt": (
            "El diagrama de Gantt es la representación visual del cronograma. Cada tarea se muestra como una "
            "barra horizontal cuya longitud indica su duración y cuya posición corresponde a su fecha de "
            "inicio. Las líneas de conexión entre barras representan las dependencias entre tareas."
        ),
        "ruta_critica": (
            "La ruta crítica es la secuencia de tareas que determina la duración mínima del proyecto. "
            "Cualquier retraso en una tarea crítica retrasa la entrega final. ProjectLibre resalta estas "
            "tareas en el diagrama de Gantt, permitiendo al gestor priorizar su seguimiento. Para el proyecto "
            "Lush Nails, la ruta crítica atraviesa el diseño de la base de datos, la autenticación, los "
            "módulos del panel y las pruebas integrales."
        ),
        "costos": (
            "Con base en la asignación de recursos y tarifas, ProjectLibre calcula el costo estimado del "
            "proyecto. Esta información es fundamental para la elaboración del presupuesto y para controlar "
            "que el proyecto se mantenga dentro del costo planificado durante su ejecución."
        ),
        "guardado": (
            "El proyecto se guarda en un archivo con extensión .pod, que almacena la planificación completa. "
            "Durante la ejecución, es posible registrar el avance real de cada tarea y compararlo con lo "
            "planificado, lo que permite dar seguimiento al proyecto."
        ),
        "aplicacion": (
            "ProjectLibre se utilizó para planificar el cronograma del proyecto Lush Nails, definiendo las "
            "tareas de desarrollo de la base de datos, el panel administrativo, el sitio web y la "
            "documentación, con sus dependencias y recursos asociados. La herramienta permitió identificar la "
            "duración total estimada del proyecto, la ruta crítica y el costo estimado, brindando un marco de "
            "control para el seguimiento del avance del sistema."
        ),
        "conclusiones": (
            "ProjectLibre constituye una herramienta de gestión de proyectos de libre acceso que permite "
            "planificar, programar y controlar los proyectos de software. Su aplicación al proyecto Portal "
            "Empresarial Lush Nails SPA permitió estructurar el trabajo en fases, definir dependencias, "
            "asignar recursos y estimar costos, facilitando la gestión del cronograma. Se recomienda mantener "
            "actualizado el registro de avance de las tareas para asegurar el control del proyecto durante su "
            "ejecución."
        ),
    }
)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(description="Genera el informe APA 7 de ProjectLibre.")
    parser.add_argument("-o", "--output", default="informe_projectlibre_apa7.docx")
    args = parser.parse_args()

    doc = Document()
    configurar_documento(doc)
    agregar_numero_pagina(doc)
    pagina_titulo(doc)
    cuerpo(doc)
    referencias_apa(doc)

    ruta_salida = os.path.join(os.path.dirname(os.path.abspath(__file__)), args.output)
    doc.save(ruta_salida)
    print(f"Informe generado: {ruta_salida}")
    print(f"Capturas encontradas: {sum(1 for k in CAPTURAS if os.path.exists(os.path.join(os.path.dirname(os.path.abspath(__file__)), 'capturas', k)))}/{len(CAPTURAS)}")


if __name__ == "__main__":
    main()
