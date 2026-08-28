#!/usr/bin/env python3
"""Genera un documento APA 7 con secciones verificadas del portal Lush Nails.

Uso: python3 informe/generar_secciones_portal_apa7.py
Salida: informe/Secciones_portal_Lush_Nails_APA7.docx
"""
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
OUT = Path(__file__).with_name("Secciones_portal_Lush_Nails_APA7.docx")


def configure(doc):
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2.54)
        section.left_margin = section.right_margin = Cm(2.54)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    header = doc.sections[0].header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    run.font.name = "Times New Roman"; run.font.size = Pt(12)


def p(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True):
    par = doc.add_paragraph(text)
    par.alignment = align
    if indent:
        par.paragraph_format.first_line_indent = Inches(.5)
    return par


def h1(doc, text):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(12)
    par.add_run(text).bold = True


def h2(doc, text):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.LEFT
    par.add_run(text).bold = True


def page_h1(doc, text):
    doc.add_page_break(); h1(doc, text)


def table(doc, number, title, headers, rows, widths=None):
    label = p(doc, f"Tabla {number}", WD_ALIGN_PARAGRAPH.LEFT, False)
    label.runs[0].bold = True
    title_p = p(doc, title, WD_ALIGN_PARAGRAPH.LEFT, False)
    title_p.runs[0].italic = True
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, value in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = value; cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs: run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            for para in cells[i].paragraphs:
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = 1
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths): row.cells[i].width = Cm(width)
    note = p(doc, "Nota. Elaboración propia a partir de la revisión del repositorio del proyecto.", WD_ALIGN_PARAGRAPH.LEFT, False)
    note.runs[0].italic = True; note.runs[0].font.size = Pt(10)


def title_page(doc):
    for _ in range(4): p(doc, "", WD_ALIGN_PARAGRAPH.CENTER, False)
    for value, bold in [
        ("Desarrollo de un Portal Empresarial Web para SPA Nails Lush", True),
        ("Secciones técnicas verificadas: puntos de función, seguridad, indicadores y planificación", True),
    ]:
        q = p(doc, value, WD_ALIGN_PARAGRAPH.CENTER, False); q.runs[0].bold = bold
    for _ in range(5): p(doc, "", WD_ALIGN_PARAGRAPH.CENTER, False)
    for value in ["Nombre del estudiante", "Nombre de la institución", "Asignatura", "Nombre del docente", "26 de agosto de 2026"]:
        p(doc, value, WD_ALIGN_PARAGRAPH.CENTER, False)
    doc.add_page_break()


def body(doc):
    h1(doc, "Resumen")
    p(doc, "Este documento complementa el informe del Portal Empresarial Web para SPA Nails Lush mediante una revisión trazable del repositorio. Se verificaron la arquitectura React–Express–PostgreSQL, el modelo de autorización basado en roles, el tablero administrativo, el Balanced Scorecard (BSC), la estructura modular y el cronograma del proyecto. Se presenta una estimación de puntos de función no ajustados, útil para dimensionamiento académico; su carácter estimado se debe a que el método IFPUG requiere la validación formal de requisitos, transacciones y tipos de datos con los usuarios. Las secciones distinguen explícitamente las funciones que el código implementa de las metas operativas propuestas para su seguimiento.")
    p(doc, "Palabras clave: portal empresarial, puntos de función, control de acceso basado en roles, Balanced Scorecard, indicadores, planificación.", indent=False)

    page_h1(doc, "Punto de Función del Portal Web")
    p(doc, "El análisis de puntos de función mide la funcionalidad entregada al usuario con independencia del lenguaje o de la tecnología. Para esta estimación se aplican pesos promedio de IFPUG: entrada externa (EI = 4), salida externa (EO = 5), consulta externa (EQ = 4), archivo lógico interno (ILF = 10) y archivo de interfaz externa (EIF = 7). El portal no consume archivos maestros externos identificables; por tanto, EIF es cero. La medición considera las capacidades persistentes y visibles en el código revisado, no el número de rutas ni de pantallas.")
    table(doc, 1, "Inventario funcional para la estimación de puntos de función", ["Tipo", "Cantidad", "Peso", "PF", "Base de conteo"], [
        ["EI", 29, 4, 116, "Registro, autenticación, citas, clientes y mantenimientos administrativos."],
        ["EO", 12, 5, 60, "Dashboard, BSC, tablero, disponibilidad, historial y listados con cálculo."],
        ["EQ", 15, 4, 60, "Consultas de servicios, sucursales, citas, galería, empleados y otros catálogos."],
        ["ILF", 19, 10, 190, "Tablas propias: sesión, catálogos, negocio, seguridad, proveedores y aliados."],
        ["EIF", 0, 7, 0, "No se identificaron archivos lógicos mantenidos por otra aplicación."],
        ["Total", "", "", 426, "Puntos de función no ajustados (PFNA)."],
    ], [1.5, 1.5, 1.5, 1.5, 10])
    p(doc, "El resultado es de 426 PFNA. No se aplica el factor de ajuste por características generales del sistema, pues no existe una valoración aprobada de comunicación de datos, rendimiento, reutilización u otros factores. En una entrega final, la cifra debe validarse con una matriz por proceso elemental y por DET/FTR, firmada por el cliente o tutor. Como evidencia técnica, el esquema inicial define 19 tablas y el backend concentra transacciones para los módulos de operación y consulta.")

    page_h1(doc, "Definición de Roles y Permisos")
    p(doc, "El portal implementa control de acceso basado en roles (RBAC). La tabla usuario_admin se relaciona con rol; rol se relaciona de manera muchos-a-muchos con permiso mediante rol_permiso. Al iniciar una petición autenticada, el middleware carga los permisos del usuario y autoriza el módulo solicitado. Las contraseñas se almacenan como hash bcrypt y la sesión se conserva mediante express-session con almacenamiento PostgreSQL.")
    table(doc, 2, "Roles implementados y alcance operativo", ["Rol", "Propósito", "Permisos asignados"], [
        ["Superadministrador", "Administración integral y control del sistema.", "Los 16 permisos del catálogo."],
        ["Administrador", "Gestión operativa completa.", "Todos excepto roles.gestionar."],
        ["Recepcionista", "Atención y agenda de clientes.", "dashboard.ver, citas.gestionar, clientes.gestionar, historial.ver y historial.gestionar."],
        ["Recursos humanos", "Gestión de talento y disponibilidad.", "dashboard.ver, empleados.gestionar y postulaciones.gestionar."],
        ["Gerencia", "Seguimiento estratégico y operativo.", "dashboard.ver, tablero.ver, bsc.ver e historial.ver."],
        ["Contabilidad", "Consulta de resultados económicos.", "tablero.ver, bsc.ver e historial.ver."],
    ], [3.2, 6.2, 6.1])
    table(doc, 3, "Catálogo de permisos por módulo", ["Módulo", "Código de permiso", "Acción autorizada"], [
        ["Dashboard", "dashboard.ver", "Consultar resumen administrativo."],
        ["Servicios", "servicios.gestionar", "Administrar servicios y categorías."],
        ["Sucursales", "sucursales.gestionar", "Administrar sedes activas."],
        ["Citas", "citas.gestionar", "Administrar agenda y estados."],
        ["Clientes", "clientes.gestionar", "Administrar información de clientes."],
        ["Empleados", "empleados.gestionar", "Administrar personal y horarios."],
        ["Galería", "galeria.gestionar", "Administrar portafolio visual."],
        ["Postulaciones", "postulaciones.gestionar", "Revisar postulaciones laborales."],
        ["Proveedores y aliados", "proveedores.gestionar / aliados.gestionar", "Administrar relaciones externas."],
        ["Configuración", "configuracion.gestionar", "Modificar información institucional."],
        ["Roles", "roles.gestionar", "Crear usuarios y administrar roles."],
        ["Analítica", "bsc.ver / tablero.ver / historial.ver", "Consultar BSC, tablero e historial."],
        ["Historial", "historial.gestionar", "Registrar servicios realizados."],
    ], [3.2, 4.5, 7.8])
    h2(doc, "Observación de seguridad")
    p(doc, "La matriz está definida en la base de datos y se consulta en el middleware; esto permite modificar permisos sin cambiar las vistas. Sin embargo, varias rutas de escritura aplican requireAuth, mientras la restricción fina depende del mapeo de rutas del middleware. Para endurecer la solución se recomienda validar el permiso también en cada operación sensible (crear, editar, eliminar), aplicar CSRF a formularios, definir cookies secure y httpOnly en producción, limitar CORS al dominio del sitio y mantener las credenciales fuera del repositorio.")

    page_h1(doc, "Identidad Individual y Modularidad")
    p(doc, "La identidad individual se expresa en dos planos. El plano institucional mantiene la marca Lush Nails SPA, su eslogan, datos de contacto, sucursales, servicios y galería. El plano de identidad digital individual identifica a cada usuario administrativo por nombre, correo electrónico único, estado activo, rol y fecha de último acceso; los clientes se registran por nombre, teléfono, correo y contraseña. Esta separación evita que las credenciales administrativas se mezclen con el acceso de clientes.")
    p(doc, "La modularidad se evidencia en una arquitectura de tres capas: React 19 para el sitio corporativo; Express con vistas EJS y API para la lógica de presentación y negocio; PostgreSQL 16 para persistencia. Las rutas del backend separan dashboard, citas, clientes, servicios, sucursales, empleados, historial, BSC, tablero, roles y entidades de apoyo. La base de datos conserva integridad mediante claves foráneas, restricciones, índices y triggers de actualización. Esta organización favorece el mantenimiento: un módulo puede evolucionar sin reescribir el portal completo, siempre que se respeten sus contratos de datos y permisos.")
    table(doc, 4, "Módulos y responsabilidad principal", ["Capa", "Módulos", "Responsabilidad"], [
        ["Sitio corporativo", "Inicio, servicios, sedes, agenda, cuenta, galería, postulación y accesibilidad.", "Atención al cliente y presencia digital."],
        ["Panel administrativo", "Dashboard, citas, clientes, empleados, horarios, historial, proveedores, aliados, roles, BSC y tablero.", "Gestión administrativa y toma de decisiones."],
        ["Persistencia", "19 tablas, vistas v_citas_completas y v_financiero, función registrar_cita y triggers.", "Consistencia, disponibilidad y trazabilidad de datos."],
    ], [3.5, 7, 5])

    page_h1(doc, "Ampliación del Dashboard")
    p(doc, "El dashboard actual consolida seis conteos operativos: servicios activos, sucursales activas, citas pendientes o confirmadas, postulaciones no leídas, proveedores activos y aliados activos. Los valores se obtienen en tiempo de consulta desde PostgreSQL, por lo que no dependen de actualización manual. Es adecuado como vista de situación inmediata para el administrador, pero no reemplaza el análisis estratégico del BSC.")
    p(doc, "La ampliación recomendada organiza el dashboard en tres niveles. El nivel operativo conserva las tarjetas actuales y agrega citas de hoy, ocupación por franja horaria y alertas de citas sin empleado asignado. El nivel táctico incorpora comparativos de 30 días, tasa de cancelación, nuevos clientes y servicios más solicitados. El nivel estratégico enlaza los indicadores consolidados del BSC: ingresos mensuales, ticket promedio, rentabilidad por servicio y cumplimiento de metas. Cada tarjeta debe mostrar periodo, fórmula, fuente y fecha/hora de actualización para prevenir interpretaciones erróneas.")
    table(doc, 5, "Propuesta de ampliación del dashboard", ["Bloque", "Indicador", "Fuente", "Uso"], [
        ["Operación diaria", "Citas de hoy y pendientes", "cita", "Priorización de agenda y confirmación."],
        ["Capacidad", "Ocupación por sucursal/empleado", "cita y horario_empleado", "Distribución de personal y horarios."],
        ["Clientes", "Nuevos clientes y cancelaciones (30 días)", "cliente y cita", "Seguimiento de demanda y retención."],
        ["Finanzas", "Ingresos mensuales y ticket promedio", "servicio_realizado", "Control de resultados económicos."],
        ["Alertas", "Postulaciones no leídas y citas sin asignar", "postulacion y cita", "Atención oportuna de pendientes."],
    ], [3.5, 4.5, 3.5, 4])

    page_h1(doc, "Indicadores y Estándares")
    p(doc, "El portal ya calcula indicadores en el BSC y en el tablero de comando. Para que funcionen como sistema de gestión, cada indicador requiere una definición estable: fórmula, fuente de datos, periodicidad, responsable, meta y regla de semáforo. Las metas incluidas en la interfaz son estándares internos configurados para el prototipo; no deben presentarse como benchmarks sectoriales sin validación de la empresa.")
    table(doc, 6, "Indicadores implementados y estándar mostrado en el portal", ["Perspectiva", "Indicador y fórmula", "Meta actual", "Fuente / frecuencia"], [
        ["Financiera", "Ingresos mensuales = SUM(monto) del mes.", "> $500", "servicio_realizado / mensual"],
        ["Financiera", "Ticket promedio = AVG(monto).", "≥ $40", "servicio_realizado / mensual"],
        ["Financiera", "Servicios activos = COUNT(servicio activo).", "≥ 30", "servicio / mensual"],
        ["Financiera", "Sucursales operativas = COUNT(sucursal activa).", "≥ 3", "sucursal / mensual"],
        ["Clientes", "Citas completadas = COUNT(estado = completada).", "≥ 10", "cita / mensual"],
        ["Clientes", "Clientes registrados = COUNT(cliente).", "≥ 10", "cliente / mensual"],
        ["Clientes", "Citas en 30 días = COUNT(fecha >= hoy − 30 días).", "≥ 5", "cita / mensual"],
        ["Procesos", "Galería activa = COUNT(galeria activa).", "≥ 5", "galeria / mensual"],
        ["Procesos", "Citas canceladas = COUNT(estado = cancelada).", "≤ 2", "cita / mensual"],
        ["Aprendizaje", "Postulaciones recibidas = COUNT(postulacion).", "≥ 1", "postulacion / mensual"],
    ], [2.4, 6.7, 2.4, 4])
    p(doc, "El tablero complementa esta medición con nuevos clientes en 30 días, atenciones del mes, tiempo promedio de atención, postulaciones pendientes, tasa de éxito de citas y citas para hoy. La tasa de éxito se calcula como citas completadas / (citas completadas + canceladas) × 100. El tiempo promedio corresponde a la diferencia entre updated_at y created_at de citas completadas; por ello mide el tiempo transcurrido registrado en el sistema y no necesariamente la duración real del servicio. Para medir duración real conviene registrar hora de inicio y finalización de la atención.")
    h2(doc, "Estándares técnicos verificables")
    p(doc, "El tablero declara como implementados: PostgreSQL 16 con relaciones normalizadas, índices y RLS; autenticación por sesión y bcrypt; comunicación API REST entre React y Express; y BSC con cuatro perspectivas. En accesibilidad, el proyecto incorpora textos alternativos, etiquetas semánticas y aria-label en los componentes; la conformidad WCAG 2.1 debe verificarse con pruebas de contraste, teclado, lector de pantalla y auditoría automatizada antes de declarar un nivel de conformidad formal (W3C, 2018).")

    page_h1(doc, "Planificación del Proyecto")
    p(doc, "La planificación existente está codificada en el archivo plan_lushnails.xml, compatible con ProjectLibre. Define inicio el 17 de agosto de 2026 y finalización el 8 de septiembre de 2026, con jornada de ocho horas, cinco días por semana. La estimación comprende 17 días laborables de calendario del proyecto. El cronograma distribuye el trabajo en cuatro paquetes: base de datos, panel administrativo, sitio corporativo y aseguramiento de calidad/documentación.")
    table(doc, 7, "Resumen de paquetes de trabajo planificados", ["Paquete", "Actividades principales", "Periodo planificado", "Responsable"], [
        ["Base de datos", "Docker/PostgreSQL, modelo E-R y vistas financieras.", "17–21 ago. 2026", "Desarrollador de BD"],
        ["Panel administrativo", "Autenticación/RBAC, citas-clientes, historial, BSC y tablero.", "18 ago.–2 sep. 2026", "Desarrollador backend"],
        ["Sitio corporativo", "React 19 e integración con API.", "24 ago.–2 sep. 2026", "Desarrollador frontend"],
        ["Calidad y entrega", "BPMN, pruebas integrales, revisión y presentación.", "20 ago.–8 sep. 2026", "Analista"],
    ], [3, 7.5, 3.5, 3])
    p(doc, "Las dependencias principales son de tipo fin-comienzo: el modelo de base de datos precede a las vistas financieras; la autenticación y roles preceden a citas y clientes; estos preceden al historial, y el historial precede al BSC/tablero. Las pruebas integrales comienzan después de completar BSC/tablero, la integración web y los diagramas BPMN; la presentación depende de las pruebas. Esta cadena debe revisarse en ProjectLibre antes de utilizarla como ruta crítica definitiva, porque cualquier cambio en duración, recurso o calendario puede modificar holguras y fecha final.")
    h2(doc, "Criterios de control")
    p(doc, "Se recomienda realizar seguimiento semanal con cuatro medidas: porcentaje de tareas completadas, variación de fechas frente a la línea base, incidencias abiertas/cerradas y evidencia de aceptación por módulo. Todo cambio en requisitos, duración o dependencia debe registrarse en el cronograma y aprobarse antes de modificar la línea base. Esta práctica permite que el diagrama de Gantt, la red crítica y el estado real reflejen la misma información de gestión.")

    page_h1(doc, "Conclusión")
    p(doc, "La revisión confirma que el Portal Empresarial Lush Nails dispone de los componentes requeridos para sustentar las secciones solicitadas: gestión de roles y permisos, módulos independientes, dashboard, BSC, tablero de comando y planificación en ProjectLibre. La estimación de 426 PFNA proporciona una referencia cuantitativa inicial del alcance funcional. Para la versión académica final, las prioridades son validar los puntos de función con requisitos aprobados, formalizar las metas de indicadores con la empresa y ejecutar pruebas de seguridad y accesibilidad que conviertan las características implementadas en evidencia verificable de calidad.")

    page_h1(doc, "Referencias")
    refs = [
        "International Function Point Users Group. (2022). Counting practices manual (Release 4.3.1). IFPUG.",
        "Lush Nails SPA. (2026). Repositorio del portal empresarial [Código fuente no publicado].",
        "OWASP Foundation. (2021). OWASP application security verification standard 4.0.3. https://owasp.org/www-project-application-security-verification-standard/",
        "Project Management Institute. (2021). A guide to the project management body of knowledge (PMBOK guide) (7th ed.).",
        "World Wide Web Consortium. (2018). Web Content Accessibility Guidelines (WCAG) 2.1. https://www.w3.org/TR/WCAG21/",
    ]
    for ref in refs:
        q = p(doc, ref, WD_ALIGN_PARAGRAPH.LEFT, False)
        q.paragraph_format.left_indent = Inches(.5); q.paragraph_format.first_line_indent = Inches(-.5)


def main():
    doc = Document(); configure(doc); title_page(doc); body(doc); doc.save(OUT)
    print(f"Documento creado: {OUT}")


if __name__ == "__main__": main()
